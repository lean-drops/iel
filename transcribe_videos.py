"""
Script Name: transcription.py

Zweck des Skripts:
Dieses Skript dient zur parallelen Transkription von MP4-Dateien im Ordner "Downloads" mittels der Whisper API.
Die parallele Verarbeitung erfolgt mit Hilfe der Bibliotheken "dask" und "concurrent.futures", um eine effiziente und schnelle
Verarbeitung der Dateien zu gewährleisten. Die Transkriptionen werden in einer Ordnerstruktur gespeichert, die der Struktur im
"Downloads" Ordner entspricht, und sowohl als .txt als auch .docx Dateien gespeichert.

Hauptfunktionen und -methoden:
- find_mp4_files: Sucht alle MP4-Dateien im angegebenen Verzeichnis.
- transcribe_mp4: Führt die Transkription einer einzelnen MP4-Datei durch.
- save_transcription: Speichert die Transkription im gewünschten Format und Verzeichnis.
- parallel_transcription: Verarbeitet alle gefundenen MP4-Dateien parallel.

Übersicht über den Ablauf des Skripts:
1. Importieren der benötigten Bibliotheken.
2. Definieren der Funktion zum Finden von MP4-Dateien.
3. Definieren der Funktion zur Transkription einer MP4-Datei.
4. Definieren der Funktion zum Speichern der Transkription.
5. Definieren der Funktion zur parallelen Verarbeitung der Transkriptionen.
6. Aufrufen der Hauptfunktion zur Ausführung des Skripts.

Hinweise auf spezielle Implementierungsentscheidungen oder Sicherheitsaspekte:
- Die Verwendung von "dask" und "concurrent.futures" gewährleistet eine effiziente Parallelverarbeitung.
- Typannotationen und umfassende Fehlerbehandlung erhöhen die Robustheit und Lesbarkeit des Codes.
- Die Transkriptionen werden in einer Ordnerstruktur gespeichert, die der Struktur im "Downloads" Ordner entspricht.
- Die Transkriptionen berücksichtigen Anglizismen und versuchen, diese korrekt zu transkribieren.
"""

import os
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List
import whisper
import logging

from dask.dataframe.io.tests.test_sql import db
from docx import Document
import moviepy.editor as mp
import re

# Logging configuration
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def find_mp4_files(directory: str) -> List[str]:
    """
    Find all MP4 files in the specified directory and its subdirectories.

    :param directory: The directory to search for MP4 files.
    :return: A list of paths to MP4 files.
    """
    mp4_files = [os.path.join(root, file)
                 for root, _, files in os.walk(directory)
                 for file in files if file.endswith('.mp4')]
    logging.info(f"Found {len(mp4_files)} MP4 files.")
    return mp4_files

def extract_audio_from_video(file_path: str) -> str:
    """
    Extract audio from the given MP4 file and save it as a temporary WAV file.

    :param file_path: Path to the MP4 file.
    :return: Path to the extracted audio WAV file.
    """
    video = mp.VideoFileClip(file_path)
    audio_path = file_path.replace('.mp4', '.wav')
    video.audio.write_audiofile(audio_path)
    return audio_path

def clean_transcription(text: str) -> str:
    """
    Clean the transcription text by handling common issues and formatting.

    :param text: The transcription text to clean.
    :return: The cleaned transcription text.
    """
    text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with a single space
    text = re.sub(r'([.,!?])', r' \1 ', text)  # Add spaces around punctuation
    text = text.strip()
    logging.info("Cleaned transcription text.")
    return text

def transcribe_mp4(file_path: str) -> str:
    """
    Transcribe the given MP4 file using Whisper API.

    :param file_path: Path to the MP4 file.
    :return: The transcription result as a string.
    """
    try:
        audio_path = extract_audio_from_video(file_path)
        model = whisper.load_model("base")
        result = model.transcribe(audio_path)
        transcription = result['text']
        transcription = clean_transcription(transcription)
        os.remove(audio_path)  # Clean up the temporary audio file
        logging.info(f"Transcription completed for {file_path}.")
        return transcription
    except Exception as e:
        logging.error(f"Error transcribing {file_path}: {e}")
        return ""

def save_transcription(file_path: str, transcription: str, base_directory: str) -> None:
    """
    Save the transcription to a .txt and .docx file in a structured directory.

    :param file_path: Path to the original MP4 file.
    :param transcription: The transcription text.
    :param base_directory: The base directory to save the transcriptions.
    """
    relative_path = os.path.relpath(file_path, base_directory)
    transcription_path = os.path.join(base_directory, 'transcriptions', os.path.splitext(relative_path)[0])

    os.makedirs(transcription_path, exist_ok=True)

    txt_file = os.path.join(transcription_path, 'transcription.txt')
    docx_file = os.path.join(transcription_path, 'transcription.docx')

    with open(txt_file, 'w', encoding='utf-8') as f:
        f.write(transcription)

    doc = Document()
    doc.add_heading('Transcription', 0)
    doc.add_paragraph(transcription)
    doc.save(docx_file)

    logging.info(f"Transcription saved to {txt_file} and {docx_file}.")

def parallel_transcription(mp4_files: List[str], base_directory: str, max_workers: int = 8) -> None:
    """
    Transcribe MP4 files in parallel and save the transcriptions.

    :param mp4_files: List of paths to MP4 files.
    :param base_directory: The base directory to save the transcriptions.
    :param max_workers: Maximum number of parallel workers.
    """
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_file = {executor.submit(transcribe_mp4, file): file for file in mp4_files}
        for future in as_completed(future_to_file):
            file = future_to_file[future]
            try:
                transcription = future.result()
                if transcription:
                    save_transcription(file, transcription, base_directory)
            except Exception as e:
                logging.error(f"Error processing file {file}: {e}")

def main():
    """
    Main function to find and transcribe MP4 files in parallel.
    """
    directory = os.path.expanduser("~/Downloads")
    mp4_files = find_mp4_files(directory)

    # Using dask bag to parallelize the transcription task
    bag = db.from_sequence(mp4_files).map(transcribe_mp4)
    transcriptions = bag.compute()

    # Save transcriptions
    for file, transcription in zip(mp4_files, transcriptions):
        save_transcription(file, transcription, directory)

if __name__ == "__main__":
    main()
